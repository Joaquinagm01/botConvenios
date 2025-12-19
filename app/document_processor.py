import os
from docx import Document
from app.config import DOCUMENT_TYPES

def fill_document_naturally(doc_type, data):
    """
    Fill document template with natural text replacement
    """
    # Map document types to template files
    template_files = {
        'niños_adolescentes': 'Convenio Niños y Adolescentes.doc',
        'tercero_directo_rcc_rcl': 'Convenio tercero directo (RCC y RCL).doc',
        'tipo_letrado_rcc_rcl': 'Convenio Tipo Letrado (RCC y RCL).doc',
        'tipo_letrado_muerte': 'Convenio Tipo Letrado Muerte.doc',
        'cesion_derechos_rcc': 'Convenio con Cesión de Derechos (RCC).doc',
        'honorarios': 'CONVENIO HONORARIOS.doc',
        'patrocinio': 'CONVENIO PATROCINIO.doc',
        'desistimiento_renuncia': 'Desistimiento Por Renuncia de Derechos.docx',
        'desistimiento_sustitucion': 'Desistimiento Sustitución de Tercero.docx',
        'recibo_pago_tercero': 'Recibo de Pago a Tercero (Efectivo).doc',
        'declaracion_no_seguro': 'Declaracion Jurada de no Seguro.doc'
    }

    template_file = template_files.get(doc_type)
    if not template_file:
        raise ValueError(f"Template not found for document type: {doc_type}")

    template_path = os.path.join('documents', template_file)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    # Load template
    doc = Document(template_path)

    # Create replacement mapping
    replacements = {
        '[NOMBRE_DEMANDANTE]': data.get('nombre_demandante', ''),
        '[DNI_DEMANDANTE]': data.get('dni_demandante', ''),
        '[DOMICILIO_DEMANDANTE]': data.get('domicilio_demandante', ''),
        '[TELEFONO_DEMANDANTE]': data.get('telefono_demandante', ''),
        '[EMAIL_DEMANDANTE]': data.get('email_demandante', ''),

        '[NOMBRE_DEMANDADO]': data.get('nombre_demandado', ''),
        '[DNI_DEMANDADO]': data.get('dni_demandado', ''),
        '[DOMICILIO_DEMANDADO]': data.get('domicilio_demandado', ''),
        '[TELEFONO_DEMANDADO]': data.get('telefono_demandado', ''),
        '[EMAIL_DEMANDADO]': data.get('email_demandado', ''),

        # Common placeholders
        '[FECHA]': '19/12/2025',  # Current date
        '[LUGAR]': 'Buenos Aires, Argentina'
    }

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Save filled document
    output_dir = 'output'
    os.makedirs(output_dir, exist_ok=True)

    output_filename = f"{doc_type}_{data.get('dni_demandante', 'unknown')}.docx"
    output_path = os.path.join(output_dir, output_filename)

    doc.save(output_path)
    return output_path